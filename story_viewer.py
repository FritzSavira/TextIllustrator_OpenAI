import os
import re
import markdown
import webbrowser
from docx import Document
from docx.shared import Cm

def create_and_open_html(directory):
    """
    Create an HTML file from text and image files in the given directory,
    open it in the default web browser, and create a corresponding DOCX file.

    Args:
        directory (str): Path to the directory containing the input files.
    """

    # Initialize HTML content with basic structure and styling
    html_content = """
    <!DOCTYPE html>
    <html lang="de">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Story View</title>
        <style>
            body { font-family: Arial, sans-serif; text-align: center; }
            h2 { font-size: 2em; }
            img { max-height: 50vh; }
            hr { border: 1px solid #000; }
            .markdown { text-align: center; display: inline-block; max-width: 80%; margin: auto; }
        </style>
    </head>
    <body>
    """

    # Add intro text if it exists
    intro_file = os.path.join(directory, "input_text.txt")
    if os.path.exists(intro_file):
        with open(intro_file, 'r', encoding='utf-8') as file:
            intro_text = file.read()
            html_content += f"<div>{intro_text}</div><hr>"

    # Filter and sort relevant files
    txt_files = sorted([f for f in os.listdir(directory) if re.match(r'^\d+scene.*\.txt$', f)])
    jpg_files = sorted([f for f in os.listdir(directory) if re.match(r'^\d+pic.*\.jpg$', f)])

    # Generate HTML content from files
    if not txt_files and not jpg_files:
        html_content += "<div>Keine passenden Dateien vorhanden.</div>"
    else:
        for txt_file in txt_files:
            scene_number = re.match(r'^(\d+)scene.*\.txt$', txt_file).group(1)
            jpg_file = next((f for f in jpg_files if f.startswith(f"{scene_number}pic")), None)

            # Add text content
            with open(os.path.join(directory, txt_file), 'r', encoding='utf-8') as file:
                scene_text = file.read()
                scene_html = markdown.markdown(scene_text)
                html_content += f'<div class="markdown">{scene_html}</div><br>'

            # Add image if it exists
            if jpg_file:
                html_content += f'<img src="{jpg_file}" alt="{jpg_file}"><hr>'

    # Close HTML structure
    html_content += """
    </body>
    </html>
    """

    # Save and open HTML file
    html_file_path = os.path.join(directory, "view_story.html")
    with open(html_file_path, 'w', encoding='utf-8') as file:
        file.write(html_content)
    webbrowser.open(html_file_path)

    # Create DOCX version
    doc = Document()
    doc.add_heading('Story View', 0)

    # Add intro to DOCX if it exists
    if os.path.exists(intro_file):
        doc.add_paragraph(intro_text)
        doc.add_paragraph('---')

    # Generate DOCX content from files
    if not txt_files and not jpg_files:
        doc.add_paragraph("Keine passenden Dateien vorhanden.")
    else:
        for txt_file in txt_files:
            scene_number = re.match(r'^(\d+)scene.*\.txt$', txt_file).group(1)
            jpg_file = next((f for f in jpg_files if f.startswith(f"{scene_number}pic")), None)

            # Add text content to DOCX
            with open(os.path.join(directory, txt_file), 'r', encoding='utf-8') as file:
                scene_text = file.read()
                scene_html = markdown.markdown(scene_text)
                doc.add_paragraph(scene_html)

            # Add image to DOCX if it exists
            if jpg_file:
                doc.add_picture(os.path.join(directory, jpg_file), width=Cm(15))
                doc.add_paragraph('---')

    # Save DOCX file
    docx_file_path = os.path.join(directory, "view_story.docx")
    doc.save(docx_file_path)